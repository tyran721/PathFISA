from __future__ import annotations

import re
import textwrap
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT / "artifacts" / "reference" / "template_code.docx"
OUTPUT_DIR = ROOT / "artifacts" / "soft-copyright"
OUTPUT_PATH = OUTPUT_DIR / "PathFISA病理小样本增量自学习智能标注软件V1.0_完整源码文档_模板格式.docx"
AUDIT_PATH = OUTPUT_DIR / "template_source_page_plan.txt"

SOFTWARE_NAME = "PathFISA病理小样本增量自学习智能标注软件"
VERSION = "V1.0"
COPYRIGHT_OWNER = "________________________"

LINES_PER_PAGE = 60
MIN_LINES_PER_PAGE = 60
MAX_LINES_PER_PAGE = 65
MAX_VISUAL_WIDTH = 58


@dataclass
class SourceLine:
    text: str
    protected: bool = False


FILE_ORDER = [
    "apps/api/app/main.py",
    "apps/api/app/imaging.py",
    "apps/api/app/models.py",
    "apps/api/app/storage.py",
    "apps/api/app/__init__.py",
    "apps/api/requirements.txt",
    "apps/web/src/main.tsx",
    "apps/web/src/api.ts",
    "apps/web/src/types.ts",
    "apps/web/src/components/AppShell.tsx",
    "apps/web/src/components/Modal.tsx",
    "apps/web/src/components/PageHeader.tsx",
    "apps/web/src/components/SlideViewer.tsx",
    "apps/web/src/pages/DashboardPage.tsx",
    "apps/web/src/pages/SlidesPage.tsx",
    "apps/web/src/pages/AnnotationPage.tsx",
    "apps/web/src/pages/TasksPage.tsx",
    "apps/web/src/pages/ModelsPage.tsx",
    "apps/web/src/pages/ModelWorkflowPage.tsx",
    "apps/web/src/pages/AnalyticsPage.tsx",
    "apps/web/src/pages/SettingsPage.tsx",
    "apps/web/src/styles.css",
    "apps/web/index.html",
    "apps/web/vite.config.ts",
    "apps/web/vite.config.d.ts",
    "apps/web/src/vite-env.d.ts",
    "apps/web/tsconfig.json",
    "apps/web/tsconfig.app.json",
    "apps/web/tsconfig.node.json",
    "apps/web/package.json",
    "apps/web/package-lock.json",
    "package.json",
    "package-lock.json",
    "scripts/setup.ps1",
]

RESERVED_LAST_PAGE = [
    "apps/web/src/App.tsx",
    "apps/web/src/components/StatusPill.tsx",
]

FILE_DESCRIPTIONS = {
    "apps/api/app/main.py": "FastAPI 后端主程序，集中实现切片库、标注、任务、模型作业、通知、智能助手、分析统计和系统设置等 API。",
    "apps/api/app/imaging.py": "病理图像读取与 Deep Zoom 支撑模块，负责 OpenSlide/Pillow 图像打开、缩略图、DZI 描述和瓦片输出。",
    "apps/api/app/models.py": "后端 Pydantic 数据模型，定义切片、标注、任务、模型运行、设置、通知和助手问答等接口结构。",
    "apps/api/app/storage.py": "本地 JSON 数据存储模块，负责标注、任务、切片登记、模型作业和项目设置等数据的读写持久化。",
    "apps/api/app/__init__.py": "后端应用包初始化文件。",
    "apps/api/requirements.txt": "后端 Python 依赖清单，声明 FastAPI、Uvicorn、OpenSlide、Pillow 等运行所需库。",
    "apps/web/src/api.ts": "前端 API 请求封装层，统一访问后端切片、标注、任务、模型、分析、通知和助手接口。",
    "apps/web/src/types.ts": "前端领域类型定义，约束切片、标注对象、任务、模型运行、设置、通知和统计数据结构。",
    "apps/web/src/components/AppShell.tsx": "应用主框架组件，实现左侧导航、页面布局、通知中心、智能助手和全局交互入口。",
    "apps/web/src/components/Modal.tsx": "通用模态对话框组件，为创建任务、导入切片、编辑配置等页面提供统一弹窗容器。",
    "apps/web/src/components/PageHeader.tsx": "页面标题与操作区组件，统一显示模块名称、说明文字、指标摘要和顶部按钮。",
    "apps/web/src/components/SlideViewer.tsx": "WSI 标注核心组件，集成 OpenSeadragon 浏览、点/矩形/多边形绘制、对象属性、AI 建议、快捷键和保存逻辑。",
    "apps/web/src/pages/DashboardPage.tsx": "项目总览页面，展示切片、标注、任务、AI 接受率、模型趋势、近期动态和系统提醒。",
    "apps/web/src/pages/SlidesPage.tsx": "病例与切片管理页面，提供真实 SVS、导入切片、切片状态、缩略图和进入标注工作台入口。",
    "apps/web/src/pages/AnnotationPage.tsx": "标注工作台页面，组合切片信息、工具栏、保存/提交操作和 SlideViewer 标注组件。",
    "apps/web/src/pages/TasksPage.tsx": "标注任务看板页面，实现任务创建、编辑、删除、筛选、状态推进和任务详情查看。",
    "apps/web/src/pages/ModelsPage.tsx": "模型中心页面，展示模型作业、评估/训练入口、运行状态、指标和模型版本记录。",
    "apps/web/src/pages/ModelWorkflowPage.tsx": "模型评估、小样本训练和增量训练配置向导，覆盖基础配置、算法架构、运行参数和质量门禁。",
    "apps/web/src/pages/AnalyticsPage.tsx": "数据分析页面，展示标注产能、标签分布、模型指标趋势、人员效率和质量统计。",
    "apps/web/src/pages/SettingsPage.tsx": "项目设置页面，维护项目基础信息、标签体系、标注流程、AI 配置、安全和通知参数。",
    "apps/web/src/styles.css": "前端全局样式表，定义医疗科研风格 UI、布局、标注工作台、看板、表单、图表和响应式样式。",
    "apps/web/index.html": "前端 HTML 入口，提供 React 应用挂载节点和页面元信息。",
    "apps/web/vite.config.ts": "Vite 构建与开发服务器配置，设置前端开发端口和 API 代理。",
    "apps/web/vite.config.d.ts": "Vite 配置的 TypeScript 声明文件。",
    "apps/web/src/vite-env.d.ts": "Vite 客户端类型声明。",
    "apps/web/tsconfig.json": "前端 TypeScript 工程主配置，组织应用和 Node 配置引用。",
    "apps/web/tsconfig.app.json": "前端应用 TypeScript 编译配置。",
    "apps/web/tsconfig.node.json": "Vite/Node 侧 TypeScript 编译配置。",
    "apps/web/package.json": "前端工程依赖与脚本配置。",
    "apps/web/package-lock.json": "前端依赖锁定清单，用于固定依赖版本和安装解析结果。",
    "package.json": "项目根级脚本配置，统一启动前端、后端、构建和检查命令。",
    "package-lock.json": "项目根级依赖锁定清单。",
    "scripts/setup.ps1": "Windows PowerShell 环境初始化脚本，用于安装前后端依赖。",
    "apps/web/src/App.tsx": "React 应用路由入口，组织 AppShell 与各业务页面路由。",
    "apps/web/src/main.tsx": "React 应用挂载入口，创建根节点并加载全局样式。",
    "apps/web/src/components/StatusPill.tsx": "状态标签组件，统一展示任务、模型和切片状态。",
}

MODULES = {
    ".py": "后端 API / 图像服务 / 数据持久化",
    ".txt": "依赖配置",
    ".ts": "前端 TypeScript 逻辑 / 构建配置",
    ".tsx": "前端 React 组件与页面",
    ".css": "前端界面样式",
    ".html": "前端页面入口",
    ".json": "工程配置 / 依赖清单",
    ".ps1": "Windows 环境初始化脚本",
}


def visual_width(text: str) -> int:
    width = 0
    for ch in text:
        width += 2 if "\u2e80" <= ch <= "\uffff" else 1
    return width


def wrap_visual(text: str, limit: int = MAX_VISUAL_WIDTH) -> list[str]:
    text = text.replace("\t", "    ").rstrip()
    if not text:
        return []
    if visual_width(text) <= limit:
        return [text]

    chunks: list[str] = []
    current = ""
    current_width = 0
    for ch in text:
        ch_width = 2 if "\u2e80" <= ch <= "\uffff" else 1
        if current and current_width + ch_width > limit:
            chunks.append(current.rstrip())
            current = ""
            current_width = 0
        current += ch
        current_width += ch_width
    if current.strip():
        chunks.append(current.rstrip())

    return chunks


def language(path: str) -> str:
    return {
        ".py": "Python",
        ".ts": "TypeScript",
        ".tsx": "TypeScript/TSX",
        ".css": "CSS",
        ".html": "HTML",
        ".json": "JSON",
        ".txt": "Text",
        ".ps1": "PowerShell",
    }.get(Path(path).suffix.lower(), "Text")


def module_name(path: str) -> str:
    return MODULES.get(Path(path).suffix.lower(), "工程源码")


def protected_block(lines: list[str]) -> list[SourceLine]:
    return [SourceLine(line, True) for line in lines]


def overall_header() -> list[SourceLine]:
    lines = [
        "/*******************************************************************************",
        f" * 软件名称：{SOFTWARE_NAME} {VERSION}",
        " * 文档名称：源程序完整清单",
        f" * 著作权人：{COPYRIGHT_OWNER}",
        " * 总体说明：本软件面向病理小样本场景，提供 Web 端智能标注与受控增量学习。",
        " * 产品定位：支持 WSI 浏览、人工标注、AI 辅助、任务复核、模型作业和统计分析。",
        " * 技术架构：前端 React + TypeScript + Vite，后端 Python FastAPI，图像浏览基于 OpenSeadragon 与 OpenSlide。",
        " * 使用边界：系统用于科研数据生产与辅助标注，不直接输出临床诊断结论。",
        " *******************************************************************************/",
    ]
    result: list[str] = []
    for line in lines:
        result.extend(wrap_visual(line))
    return protected_block(result)


def file_header(path: str) -> list[SourceLine]:
    desc = FILE_DESCRIPTIONS.get(path, "项目源程序文件。")
    raw_lines = [
        "/*******************************************************************************",
        f" * 软件名称：{SOFTWARE_NAME} {VERSION}",
        f" * 源文件：{path}",
        f" * 编程语言/类型：{language(path)}",
        f" * 所属模块：{module_name(path)}",
        f" * 功能说明：{desc}",
        " * 文档说明：以下为该文件源代码，按申请材料格式连续排版。",
        " *******************************************************************************/",
    ]
    lines: list[str] = []
    for line in raw_lines:
        lines.extend(wrap_visual(line))
    return protected_block(lines)


def source_file_lines(path: str) -> list[SourceLine]:
    lines = file_header(path)
    source = ROOT / path
    raw_text = source.read_text(encoding="utf-8", errors="replace")
    raw_lines = raw_text.splitlines()
    if not raw_lines or not any(line.strip() for line in raw_lines):
        lines.append(SourceLine("/* 本文件为空包初始化文件，无可执行语句。 */"))
        return lines
    for raw in raw_lines:
        if not raw.strip():
            continue
        for wrapped in wrap_visual(raw):
            lines.append(SourceLine(wrapped))
    return lines


def collect_lines(paths: list[str]) -> list[SourceLine]:
    lines: list[SourceLine] = []
    for path in paths:
        lines.extend(source_file_lines(path))
    return lines


def inside_protected(lines: list[SourceLine], index: int) -> bool:
    return 0 < index < len(lines) and lines[index].protected and lines[index - 1].protected


def protected_start(lines: list[SourceLine], index: int) -> int:
    while index > 0 and lines[index - 1].protected:
        index -= 1
    return index


def protected_end(lines: list[SourceLine], index: int) -> int:
    while index < len(lines) and lines[index].protected:
        index += 1
    return index


def split_pages(lines: list[SourceLine]) -> list[list[SourceLine]]:
    pages: list[list[SourceLine]] = []
    start = 0
    total = len(lines)
    while start < total:
        remaining = total - start
        if remaining <= LINES_PER_PAGE:
            if remaining < MIN_LINES_PER_PAGE and pages:
                need = MIN_LINES_PER_PAGE - remaining
                move = min(need, len(pages[-1]) - MIN_LINES_PER_PAGE)
                if move > 0:
                    tail = pages[-1][-move:]
                    pages[-1] = pages[-1][:-move]
                    pages.append(tail + lines[start:])
                    break
            pages.append(lines[start:])
            break
        split = start + LINES_PER_PAGE
        if inside_protected(lines, split):
            before = protected_start(lines, split)
            after = protected_end(lines, split)
            if before - start >= MIN_LINES_PER_PAGE:
                split = before
            elif after - start <= LINES_PER_PAGE:
                split = after
        pages.append(lines[start:split])
        start = split
    return rebalance_tail(pages)


def rebalance_tail(pages: list[list[SourceLine]]) -> list[list[SourceLine]]:
    if not pages or len(pages[-1]) >= MIN_LINES_PER_PAGE:
        return pages

    for suffix_count in range(2, min(30, len(pages)) + 1):
        combined = [line for page in pages[-suffix_count:] for line in page]
        total = len(combined)
        if total < suffix_count * MIN_LINES_PER_PAGE:
            reduced_count = suffix_count - 1
            if reduced_count < 1:
                continue
            if not (reduced_count * MIN_LINES_PER_PAGE <= total <= reduced_count * MAX_LINES_PER_PAGE):
                continue
            balanced = split_evenly(combined, reduced_count)
            if balanced and all(MIN_LINES_PER_PAGE <= len(page) <= MAX_LINES_PER_PAGE for page in balanced):
                return pages[:-suffix_count] + balanced
            continue

        balanced = split_evenly(combined, suffix_count)
        if balanced and all(MIN_LINES_PER_PAGE <= len(page) <= MAX_LINES_PER_PAGE for page in balanced):
            return pages[:-suffix_count] + balanced

    return pages


def split_evenly(lines: list[SourceLine], page_count: int) -> list[list[SourceLine]] | None:
    boundaries = [0]
    total = len(lines)
    for page_index in range(1, page_count):
        ideal = round(total * page_index / page_count)
        lower = boundaries[-1] + MIN_LINES_PER_PAGE
        upper = total - (page_count - page_index) * MIN_LINES_PER_PAGE
        split = max(lower, min(ideal, upper))
        if inside_protected(lines, split):
            left = protected_start(lines, split)
            right = protected_end(lines, split)
            if left >= lower:
                split = left
            elif right <= upper:
                split = right
            else:
                return None
        boundaries.append(split)
    boundaries.append(total)
    return [lines[boundaries[i] : boundaries[i + 1]] for i in range(page_count)]


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_or_append(parent, element_name: str) -> OxmlElement:
    node = parent.find(qn(element_name))
    if node is None:
        node = OxmlElement(element_name)
        parent.append(node)
    return node


def adjust_template_line_density(doc: Document) -> None:
    """Keep the user's template geometry/font and only tighten the line grid enough for 60+ lines/page."""
    normal = doc.styles["Normal"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(10.6)

    ppr = normal._element.get_or_add_pPr()
    snap = set_or_append(ppr, "w:snapToGrid")
    snap.set(qn("w:val"), "0")
    spacing = set_or_append(ppr, "w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "212")
    spacing.set(qn("w:lineRule"), "exact")

    for section in doc.sections:
        sect_pr = section._sectPr
        doc_grid = sect_pr.find(qn("w:docGrid"))
        if doc_grid is not None:
            doc_grid.set(qn("w:linePitch"), "212")


def add_source_line(doc: Document, line: SourceLine) -> None:
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(10.6)
    p.paragraph_format.widow_control = False
    p.add_run(line.text)


def build_document() -> None:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")
    missing = [path for path in FILE_ORDER + RESERVED_LAST_PAGE if not (ROOT / path).exists()]
    if missing:
        raise FileNotFoundError("Missing source files: " + ", ".join(missing))

    main_lines = overall_header() + collect_lines(FILE_ORDER)
    reserved_last = collect_lines(RESERVED_LAST_PAGE)
    pages = split_pages(main_lines) + [reserved_last]
    for index, page in enumerate(pages, start=1):
        if len(page) < MIN_LINES_PER_PAGE:
            raise RuntimeError(f"Page {index} has only {len(page)} planned lines.")

    doc = Document(str(TEMPLATE_PATH))
    clear_document_body(doc)
    adjust_template_line_density(doc)
    doc.core_properties.title = f"{SOFTWARE_NAME}{VERSION}完整源码文档"
    doc.core_properties.subject = "计算机软件著作权登记源程序完整清单"
    doc.core_properties.author = "PathFISA"

    for page_index, page in enumerate(pages, start=1):
        for line in page:
            add_source_line(doc, line)
        if page_index < len(pages):
            doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    AUDIT_PATH.write_text(
        "\n".join(
            [f"page={i:03d}, planned_lines={len(page)}" for i, page in enumerate(pages, start=1)]
            + [f"planned_pages={len(pages)}"]
        ),
        encoding="utf-8",
    )


if __name__ == "__main__":
    build_document()
    print(OUTPUT_PATH)
