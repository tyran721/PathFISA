from __future__ import annotations

import math
import re
import textwrap
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "artifacts" / "soft-copyright"
OUTPUT_PATH = OUTPUT_DIR / "PathFISA病理小样本增量自学习智能标注软件V1.0_完整源码文档.docx"
AUDIT_PATH = OUTPUT_DIR / "full_source_page_audit.txt"

SOFTWARE_NAME = "PathFISA病理小样本增量自学习智能标注软件"
VERSION = "V1.0"
COPYRIGHT_OWNER = "________________________"

INK = "111111"
HEADER = "17365D"
MUTED = "555555"

TARGET_LINES_PER_PAGE = 72
MIN_LINES_PER_PAGE = 60
MAX_LINES_PER_PAGE = 84
WRAP_WIDTH = 158


@dataclass
class DisplayLine:
    text: str
    kind: str = "code"
    protected: bool = False


FILE_ORDER = [
    "apps/api/app/main.py",
    "apps/api/app/imaging.py",
    "apps/api/app/models.py",
    "apps/api/app/storage.py",
    "apps/api/app/__init__.py",
    "apps/api/requirements.txt",
    "apps/web/src/api.ts",
    "apps/web/src/types.ts",
    "apps/web/src/components/AppShell.tsx",
    "apps/web/src/components/Modal.tsx",
    "apps/web/src/components/PageHeader.tsx",
    "apps/web/src/components/SlideViewer.tsx",
    "apps/web/src/pages/DashboardPage.tsx",
    "apps/web/src/pages/SlidesPage.tsx",
    "apps/web/src/pages/AnnotationPage.tsx",
    "apps/web/src/pages/TasksPage.tsx",
    "apps/web/src/pages/ModelsPage.tsx",
    "apps/web/src/pages/ModelWorkflowPage.tsx",
    "apps/web/src/pages/AnalyticsPage.tsx",
    "apps/web/src/pages/SettingsPage.tsx",
    "apps/web/src/styles.css",
    "apps/web/index.html",
    "apps/web/vite.config.ts",
    "apps/web/vite.config.d.ts",
    "apps/web/src/vite-env.d.ts",
    "apps/web/tsconfig.json",
    "apps/web/tsconfig.app.json",
    "apps/web/tsconfig.node.json",
    "apps/web/package.json",
    "apps/web/package-lock.json",
    "package.json",
    "package-lock.json",
    "scripts/setup.ps1",
]

RESERVED_LAST_PAGE = [
    "apps/web/src/App.tsx",
    "apps/web/src/main.tsx",
    "apps/web/src/components/StatusPill.tsx",
]

FILE_DESCRIPTIONS = {
    "apps/api/app/main.py": "FastAPI 后端主程序，集中实现切片库、标注、任务、模型作业、通知、智能助手、分析统计和系统设置等 API。",
    "apps/api/app/imaging.py": "病理图像读取与 Deep Zoom 支撑模块，负责 OpenSlide/Pillow 图像打开、缩略图、DZI 描述和瓦片输出。",
    "apps/api/app/models.py": "后端 Pydantic 数据模型，定义切片、标注、任务、模型运行、设置、通知和助手问答等接口结构。",
    "apps/api/app/storage.py": "本地 JSON 数据存储模块，负责标注、任务、切片登记、模型作业和项目设置等数据的读写持久化。",
    "apps/api/requirements.txt": "后端 Python 依赖清单，声明 FastAPI、Uvicorn、OpenSlide、Pillow 等运行所需库。",
    "apps/web/src/api.ts": "前端 API 请求封装层，统一访问后端切片、标注、任务、模型、分析、通知和助手接口。",
    "apps/web/src/types.ts": "前端领域类型定义，约束切片、标注对象、任务、模型运行、设置、通知和统计数据结构。",
    "apps/web/src/components/AppShell.tsx": "应用主框架组件，实现左侧导航、页面布局、通知中心、智能助手和全局交互入口。",
    "apps/web/src/components/Modal.tsx": "通用模态对话框组件，为创建任务、导入切片、编辑配置等页面提供统一弹窗容器。",
    "apps/web/src/components/PageHeader.tsx": "页面标题与操作区组件，统一显示模块名称、说明文字、指标摘要和顶部按钮。",
    "apps/web/src/components/SlideViewer.tsx": "WSI 标注核心组件，集成 OpenSeadragon 浏览、点/矩形/多边形绘制、对象属性、AI 建议、快捷键和保存逻辑。",
    "apps/web/src/pages/DashboardPage.tsx": "项目总览页面，展示切片、标注、任务、AI 接受率、模型趋势、近期动态和系统提醒。",
    "apps/web/src/pages/SlidesPage.tsx": "病例与切片管理页面，提供真实 SVS、导入切片、切片状态、缩略图和进入标注工作台入口。",
    "apps/web/src/pages/AnnotationPage.tsx": "标注工作台页面，组合切片信息、工具栏、保存/提交操作和 SlideViewer 标注组件。",
    "apps/web/src/pages/TasksPage.tsx": "标注任务看板页面，实现任务创建、编辑、删除、筛选、状态推进和任务详情查看。",
    "apps/web/src/pages/ModelsPage.tsx": "模型中心页面，展示模型作业、评估/训练入口、运行状态、指标和模型版本记录。",
    "apps/web/src/pages/ModelWorkflowPage.tsx": "模型评估、小样本训练和增量训练配置向导，覆盖基础配置、算法架构、运行参数和质量门禁。",
    "apps/web/src/pages/AnalyticsPage.tsx": "数据分析页面，展示标注产能、标签分布、模型指标趋势、人员效率和质量统计。",
    "apps/web/src/pages/SettingsPage.tsx": "项目设置页面，维护项目基础信息、标签体系、标注流程、AI 配置、安全和通知参数。",
    "apps/web/src/styles.css": "前端全局样式表，定义医疗科研风格 UI、布局、标注工作台、看板、表单、图表和响应式样式。",
    "apps/web/index.html": "前端 HTML 入口，提供 React 应用挂载节点和页面元信息。",
    "apps/web/vite.config.ts": "Vite 构建与开发服务器配置，设置前端开发端口和 API 代理。",
    "apps/web/tsconfig.json": "前端 TypeScript 工程主配置，组织应用和 Node 配置引用。",
    "apps/web/tsconfig.app.json": "前端应用 TypeScript 编译配置。",
    "apps/web/tsconfig.node.json": "Vite/Node 侧 TypeScript 编译配置。",
    "apps/web/package.json": "前端工程依赖与脚本配置。",
    "apps/web/package-lock.json": "前端依赖锁定清单，用于固定依赖版本和安装解析结果。",
    "package.json": "项目根级脚本配置，统一启动前端、后端、构建和检查命令。",
    "package-lock.json": "项目根级依赖锁定清单。",
    "scripts/setup.ps1": "Windows PowerShell 环境初始化脚本，用于安装前后端依赖。",
    "apps/web/src/App.tsx": "React 应用路由入口，组织 AppShell 与各业务页面路由。",
    "apps/web/src/main.tsx": "React 应用挂载入口，创建根节点并加载全局样式。",
    "apps/web/src/components/StatusPill.tsx": "状态标签组件，统一展示任务、模型和切片状态。",
    "apps/web/src/vite-env.d.ts": "Vite 客户端类型声明。",
    "apps/web/vite.config.d.ts": "Vite 配置的 TypeScript 声明文件。",
    "apps/api/app/__init__.py": "后端应用包初始化文件。",
}

MODULE_MAP = {
    ".py": "后端 API / 图像服务 / 数据持久化",
    ".txt": "依赖配置",
    ".ts": "前端 TypeScript 逻辑 / 构建配置",
    ".tsx": "前端 React 组件与页面",
    ".css": "前端界面样式",
    ".html": "前端页面入口",
    ".json": "工程配置 / 依赖清单",
    ".ps1": "Windows 环境初始化脚本",
}


def set_run_font(run, *, size: float, color: str = INK, bold: bool | None = None) -> None:
    run.font.name = "Consolas"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Consolas")
    rpr.rFonts.set(qn("w:hAnsi"), "Consolas")
    rpr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def language_for_path(path: str) -> str:
    suffix = Path(path).suffix.lower()
    return {
        ".py": "Python",
        ".ts": "TypeScript",
        ".tsx": "TypeScript/TSX",
        ".css": "CSS",
        ".html": "HTML",
        ".json": "JSON",
        ".txt": "Text",
        ".ps1": "PowerShell",
    }.get(suffix, "Text")


def module_for_path(path: str) -> str:
    return MODULE_MAP.get(Path(path).suffix.lower(), "工程源码")


def clean_text(text: str) -> str:
    return text.replace("\t", "    ").rstrip("\r\n")


def wrap_code_line(text: str) -> list[str]:
    if not text:
        return []
    expanded = clean_text(text)
    if len(expanded) <= WRAP_WIDTH:
        return [expanded]

    indent = re.match(r"^\s*", expanded).group(0)
    wrapper = textwrap.TextWrapper(
        width=WRAP_WIDTH,
        replace_whitespace=False,
        drop_whitespace=False,
        break_long_words=True,
        break_on_hyphens=False,
        subsequent_indent=indent + "    ",
    )
    return [part.rstrip() for part in wrapper.wrap(expanded) if part.strip()]


def block(lines: list[str], kind: str = "header") -> list[DisplayLine]:
    return [DisplayLine(text=line, kind=kind, protected=True) for line in lines]


def overall_header() -> list[DisplayLine]:
    lines = [
        "/*******************************************************************************",
        f" * 软件名称：{SOFTWARE_NAME} {VERSION}",
        " * 文档名称：源程序完整清单",
        f" * 著作权人：{COPYRIGHT_OWNER}",
        " * 总体说明：本软件是面向病理小样本场景的 Web 端智能标注与受控增量学习平台。",
        " * 产品定位：支持 WSI 浏览、人工标注、AI 辅助标注、任务复核、模型作业、分析统计与项目设置。",
        " * 技术架构：前端采用 React + TypeScript + Vite，后端采用 Python FastAPI，图像浏览基于 OpenSeadragon 与 OpenSlide。",
        " * 业务边界：系统用于科研数据生产与辅助标注，不直接输出临床诊断结论；模型更新须经人工复核和发布控制。",
        " * 说明来源：本段依据 README、需求分析和系统设计文档整理，用于源码文档功能说明。",
        " *******************************************************************************/",
    ]
    return block(lines)


def file_header(path: str) -> list[DisplayLine]:
    description = FILE_DESCRIPTIONS.get(path, "项目源程序文件。")
    lines = [
        "/*******************************************************************************",
        f" * 软件名称：{SOFTWARE_NAME} {VERSION}",
        f" * 源文件：{path}",
        f" * 编程语言/类型：{language_for_path(path)}",
        f" * 所属模块：{module_for_path(path)}",
        f" * 功能说明：{description}",
        " * 文档说明：以下内容为该文件源码逐行摘录，保留主要业务逻辑、配置和界面实现。",
        " *******************************************************************************/",
    ]
    return block(lines)


def read_source_lines(path: str) -> list[DisplayLine]:
    source_path = ROOT / path
    lines = file_header(path)
    raw_text = source_path.read_text(encoding="utf-8", errors="replace")
    raw_lines = raw_text.splitlines()
    if not raw_lines or not any(line.strip() for line in raw_lines):
        lines.append(DisplayLine("/* 本文件为空包初始化文件，无可执行语句。 */", kind="code"))
        return lines

    for raw in raw_lines:
        if not raw.strip():
            continue
        for piece in wrap_code_line(raw):
            lines.append(DisplayLine(piece, kind="code"))
    return lines


def build_display_lines(paths: list[str]) -> list[DisplayLine]:
    result: list[DisplayLine] = []
    for path in paths:
        result.extend(read_source_lines(path))
    return result


def is_inside_header(lines: list[DisplayLine], index: int) -> bool:
    if index <= 0 or index >= len(lines):
        return False
    return lines[index].protected and lines[index - 1].protected


def header_start(lines: list[DisplayLine], index: int) -> int:
    while index > 0 and lines[index - 1].protected:
        index -= 1
    return index


def header_end(lines: list[DisplayLine], index: int) -> int:
    while index < len(lines) and lines[index].protected:
        index += 1
    return index


def partition_main_lines(lines: list[DisplayLine]) -> list[list[DisplayLine]]:
    pages: list[list[DisplayLine]] = []
    start = 0
    total = len(lines)
    while start < total:
        remaining = total - start
        if remaining <= MAX_LINES_PER_PAGE:
            if remaining < MIN_LINES_PER_PAGE and pages:
                need = MIN_LINES_PER_PAGE - remaining
                move = min(need, len(pages[-1]) - MIN_LINES_PER_PAGE)
                if move > 0:
                    tail = pages[-1][-move:]
                    pages[-1] = pages[-1][:-move]
                    pages.append(tail + lines[start:])
                    break
            pages.append(lines[start:])
            break

        split = min(start + TARGET_LINES_PER_PAGE, total)
        if is_inside_header(lines, split):
            before_header = header_start(lines, split)
            after_header = header_end(lines, split)
            if before_header - start >= MIN_LINES_PER_PAGE:
                split = before_header
            elif after_header - start <= MAX_LINES_PER_PAGE:
                split = after_header

        if split - start < MIN_LINES_PER_PAGE:
            split = min(start + MIN_LINES_PER_PAGE, total)
            if is_inside_header(lines, split):
                split = header_end(lines, split)
        if split - start > MAX_LINES_PER_PAGE:
            split = start + MAX_LINES_PER_PAGE
            if is_inside_header(lines, split):
                split = header_start(lines, split)
        pages.append(lines[start:split])
        start = split
    return rebalance_underfilled_tail(pages)


def split_evenly_with_safe_headers(lines: list[DisplayLine], page_count: int) -> list[list[DisplayLine]]:
    boundaries = [0]
    total = len(lines)
    for page_index in range(1, page_count):
        ideal = round(total * page_index / page_count)
        lower = boundaries[-1] + MIN_LINES_PER_PAGE
        upper = total - (page_count - page_index) * MIN_LINES_PER_PAGE
        split = max(lower, min(ideal, upper))
        if is_inside_header(lines, split):
            before = header_start(lines, split)
            after = header_end(lines, split)
            if before >= lower:
                split = before
            elif after <= upper:
                split = after
        boundaries.append(split)
    boundaries.append(total)
    return [lines[boundaries[i] : boundaries[i + 1]] for i in range(page_count)]


def rebalance_underfilled_tail(pages: list[list[DisplayLine]]) -> list[list[DisplayLine]]:
    if not pages or len(pages[-1]) >= MIN_LINES_PER_PAGE:
        return pages

    for suffix_count in range(2, min(10, len(pages)) + 1):
        suffix_pages = pages[-suffix_count:]
        suffix_lines = [line for page in suffix_pages for line in page]
        total = len(suffix_lines)
        if MIN_LINES_PER_PAGE * suffix_count <= total <= MAX_LINES_PER_PAGE * suffix_count:
            balanced = split_evenly_with_safe_headers(suffix_lines, suffix_count)
            if all(MIN_LINES_PER_PAGE <= len(page) <= MAX_LINES_PER_PAGE for page in balanced):
                return pages[:-suffix_count] + balanced

    return pages


def balance_pages(pages: list[list[DisplayLine]], reserved_last: list[DisplayLine]) -> list[list[DisplayLine]]:
    all_pages = pages + [reserved_last]
    for index, page in enumerate(all_pages, start=1):
        if len(page) < MIN_LINES_PER_PAGE:
            raise RuntimeError(f"Page {index} has {len(page)} display lines, expected at least {MIN_LINES_PER_PAGE}.")
        if len(page) > MAX_LINES_PER_PAGE + 8:
            raise RuntimeError(f"Page {index} has {len(page)} display lines and may overflow.")
    return all_pages


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.15)
    section.right_margin = Cm(1.0)
    section.header_distance = Cm(0.35)
    section.footer_distance = Cm(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Consolas"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(5.9)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(7.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_before = Pt(0)
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run(f"{SOFTWARE_NAME} {VERSION} 源程序完整清单")
    set_run_font(run, size=7.2, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    run = footer.add_run("第 ")
    set_run_font(run, size=7.2, color=MUTED)
    add_field(footer, "PAGE")
    run = footer.add_run(" 页")
    set_run_font(run, size=7.2, color=MUTED)

    doc.core_properties.title = f"{SOFTWARE_NAME}{VERSION}完整源码文档"
    doc.core_properties.subject = "计算机软件著作权登记源程序完整清单"
    doc.core_properties.author = "PathFISA"
    doc.core_properties.keywords = "PathFISA, 源程序, 软件著作权, Python, TypeScript, WSI"


def add_display_line(doc: Document, line: DisplayLine) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(7.5)
    paragraph.paragraph_format.keep_together = False
    paragraph.paragraph_format.keep_with_next = False
    paragraph.paragraph_format.widow_control = False
    run = paragraph.add_run(line.text)
    if line.kind == "header":
        set_run_font(run, size=5.9, color=HEADER, bold=True)
    else:
        set_run_font(run, size=5.9, color=INK)


def build_document() -> None:
    missing = [path for path in FILE_ORDER + RESERVED_LAST_PAGE if not (ROOT / path).exists()]
    if missing:
        raise FileNotFoundError("Missing source files: " + ", ".join(missing))

    main_lines = overall_header() + build_display_lines(FILE_ORDER)
    reserved_last = build_display_lines(RESERVED_LAST_PAGE)
    pages = balance_pages(partition_main_lines(main_lines), reserved_last)

    doc = Document()
    configure_document(doc)
    for page_index, page in enumerate(pages, start=1):
        for line in page:
            add_display_line(doc, line)
        if page_index < len(pages):
            doc.paragraphs[-1].runs[-1].add_break(WD_BREAK.PAGE)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)

    audit = [
        f"page={index:03d}, display_lines={len(page)}, header_start={page[0].text.startswith('/*****')}"
        for index, page in enumerate(pages, start=1)
    ]
    audit.append(f"total_pages_planned={len(pages)}")
    audit.append(f"total_display_lines={sum(len(page) for page in pages)}")
    audit.append(f"included_files={len(FILE_ORDER) + len(RESERVED_LAST_PAGE)}")
    AUDIT_PATH.write_text("\n".join(audit), encoding="utf-8")


if __name__ == "__main__":
    build_document()
    print(OUTPUT_PATH)
