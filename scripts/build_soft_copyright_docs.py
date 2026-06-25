from __future__ import annotations

import math
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "artifacts" / "soft-copyright"
MANUAL_PATH = OUTPUT_DIR / "PathFISA病理小样本增量自学习智能标注软件V1.0_用户手册.docx"
SOURCE_PATH = OUTPUT_DIR / "PathFISA病理小样本增量自学习智能标注软件V1.0_源码文档.docx"

SOFTWARE_NAME = "PathFISA病理小样本增量自学习智能标注软件"
VERSION = "V1.0"
COPYRIGHT_OWNER = "________________________"
PREPARED_BY = "________________________"

NAVY = "12372A"
GREEN = "18966B"
MINT = "E8F5EF"
LIGHT = "F3F7F5"
GRAY = "66736D"
INK = "17221E"
LINE = "D7E4DE"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6, style="single") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), style)
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_widths(table, widths_dxa: list[int], indent_dxa=120) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, latin="Arial", east_asia="微软雅黑", size=None, color=None, bold=None, italic=None):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def set_manual_page_geometry(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.75)
    section.different_first_page_header_footer = True


def configure_manual_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after, color in (
        ("Title", 27, 0, 8, NAVY),
        ("Subtitle", 13, 0, 8, GRAY),
        ("Heading 1", 17, 14, 8, NAVY),
        ("Heading 2", 13, 10, 5, GREEN),
        ("Heading 3", 11, 8, 4, NAVY),
    ):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Cm(0.7)
        style.paragraph_format.first_line_indent = Cm(-0.35)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2


def add_manual_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{SOFTWARE_NAME} {VERSION} · 用户手册")
    set_run_font(run, size=8.5, color=GRAY, bold=True)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), LINE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("第 ")
    set_run_font(r, size=8.5, color=GRAY)
    add_field(fp, "PAGE")
    r = fp.add_run(" 页 / 共 ")
    set_run_font(r, size=8.5, color=GRAY)
    add_field(fp, "NUMPAGES")
    r = fp.add_run(" 页")
    set_run_font(r, size=8.5, color=GRAY)


def add_kicker(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text.upper())
    set_run_font(run, size=8, color=GREEN, bold=True)


def add_page_title(doc: Document, chapter: str, title: str, lead: str | None = None) -> None:
    add_kicker(doc, chapter)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.add_run(title)
    if lead:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_after = Pt(8)
        run = lp.add_run(lead)
        set_run_font(run, size=10, color=GRAY)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def add_callout(doc: Document, label: str, text: str, fill=MINT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [9360], indent_dxa=120)
    set_table_borders(table, color=fill, size=0, style="nil")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=110, start=150, bottom=110, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{label}：")
    set_run_font(run, size=9.5, color=GREEN, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=9.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_picture(doc: Document, filename: str, width_cm=16.7, caption: str | None = None) -> None:
    path = ROOT / "artifacts" / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(6)
        r = cp.add_run(caption)
        set_run_font(r, size=8.5, color=GRAY)


def add_two_pictures(doc: Document, left: tuple[str, str], right: tuple[str, str]) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [4680, 4680], indent_dxa=0)
    set_table_borders(table, color=WHITE, size=0, style="nil")
    for cell, (filename, caption) in zip(table.rows[0].cells, (left, right)):
        set_cell_margins(cell, top=40, start=40, bottom=40, end=40)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(ROOT / "artifacts" / filename), width=Cm(7.9))
        cp = cell.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(0)
        r = cp.add_run(caption)
        set_run_font(r, size=8, color=GRAY)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def add_metadata_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [2200, 7160], indent_dxa=120)
    set_table_borders(table, color=LINE, size=5)
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], LIGHT)
        for cell in cells:
            set_cell_margins(cell, top=100, start=130, bottom=100, end=130)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cells[0].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=9.5, color=INK)


def build_manual() -> None:
    doc = Document()
    set_manual_page_geometry(doc)
    configure_manual_styles(doc)
    add_manual_header_footer(doc)
    doc.core_properties.title = f"{SOFTWARE_NAME}{VERSION}用户手册"
    doc.core_properties.subject = "计算机软件著作权登记配套用户手册"
    doc.core_properties.author = "PathFISA"
    doc.core_properties.keywords = "PathFISA, 病理标注, WSI, 用户手册, 软件著作权"

    # Page 1: cover (editorial_cover pattern, A4 named override).
    doc.add_paragraph().paragraph_format.space_after = Pt(74)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("软件著作权登记材料")
    set_run_font(r, size=10, color=GREEN, bold=True)
    p.paragraph_format.space_after = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(SOFTWARE_NAME)
    set_run_font(r, size=26, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Pathology Few-shot Incremental Self-learning Annotation")
    set_run_font(r, size=11, color=GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(42)
    r = p.add_run(f"{VERSION} 用户手册")
    set_run_font(r, size=18, color=GREEN, bold=True)

    add_metadata_table(
        doc,
        [
            ("软件版本", VERSION),
            ("著作权人", COPYRIGHT_OWNER),
            ("编制单位", PREPARED_BY),
            ("编制日期", "2026年06月"),
            ("文档用途", "计算机软件著作权登记及软件使用说明"),
        ],
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(40)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本软件用于科研数据生产与辅助标注，不直接输出临床诊断结论。")
    set_run_font(r, size=9, color=GRAY, italic=True)
    page_break(doc)

    # Page 2
    add_page_title(doc, "文档信息", "文档说明", "本手册依据当前可运行版本的界面、接口和工程结构编写。")
    add_metadata_table(
        doc,
        [
            ("文档名称", f"{SOFTWARE_NAME}{VERSION}用户手册"),
            ("软件简称", "PathFISA"),
            ("版本号", VERSION),
            ("运行形态", "浏览器端 Web 应用，前后端分离"),
            ("适用对象", "病理医生、标注人员、项目管理员、算法工程师"),
            ("主要用途", "WSI 浏览、人工/AI 辅助标注、任务管理、模型作业与数据分析"),
        ],
    )
    doc.add_heading("版本记录", level=2)
    table = doc.add_table(rows=2, cols=4)
    set_table_widths(table, [1300, 1800, 2360, 3900], indent_dxa=120)
    set_table_borders(table)
    headers = ["版本", "日期", "编制", "说明"]
    for index, text in enumerate(headers):
        set_cell_shading(table.rows[0].cells[index], MINT)
        p = table.rows[0].cells[index].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=9, color=NAVY, bold=True)
    values = [VERSION, "2026-06", PREPARED_BY, "首次形成软著登记用户手册"]
    for index, text in enumerate(values):
        p = table.rows[1].cells[index].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index < 3 else WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(text)
        set_run_font(r, size=9, color=INK)
    add_callout(doc, "说明", "界面中的示例病例、指标和模型结果用于演示系统流程；实际使用时以机构导入的数据和配置为准。")
    page_break(doc)

    # Page 3
    add_page_title(doc, "目录", "内容导航")
    toc_items = [
        ("1", "软件概述", "4"),
        ("2", "运行环境与安装", "5"),
        ("3", "启动与访问", "6"),
        ("4", "主界面与导航", "7"),
        ("5", "通知中心与智能助手", "8"),
        ("6", "病例与切片管理", "9"),
        ("7", "切片导入", "10"),
        ("8", "标注工作台", "11"),
        ("9", "浏览与显示参数", "12"),
        ("10", "创建标注与快捷键", "13"),
        ("11", "标注属性与 AI 建议", "14"),
        ("12", "保存、提交与复核", "15"),
        ("13", "标注任务管理", "16"),
        ("14", "模型中心与模型作业", "17"),
        ("15", "增量训练配置", "18"),
        ("16", "数据分析", "19"),
        ("17", "项目设置与数据管理", "20"),
        ("18", "常见问题", "21"),
        ("19", "安全合规与使用边界", "22"),
        ("附录", "接口、数据目录与验收清单", "23"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_widths(table, [1000, 7460, 900], indent_dxa=120)
    set_table_borders(table, color=LINE, size=4)
    for idx, text in enumerate(("章节", "名称", "页码")):
        set_cell_shading(table.rows[0].cells[idx], MINT)
        p = table.rows[0].cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    for number, name, page in toc_items:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_margins(cell, top=65, start=110, bottom=65, end=110)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell, text in zip(cells, (number, name, page)):
            r = cell.paragraphs[0].add_run(text)
            set_run_font(r, size=9, color=INK)
    page_break(doc)

    # Page 4
    add_page_title(doc, "第1章", "软件概述", "PathFISA 面向病理小样本场景，提供从切片管理到受控模型迭代的一体化工作空间。")
    doc.add_heading("1.1 软件定位", level=2)
    doc.add_paragraph(
        "本软件以全切片图像（Whole Slide Image，WSI）为主要对象，通过浏览器提供多分辨率浏览、"
        "人工标注、AI 预标注复核、任务协作、模型评估、小样本训练、增量训练和数据分析能力。"
    )
    doc.add_heading("1.2 主要功能", level=2)
    for item in (
        "切片库：管理 SVS、TIFF、NDPI、PNG、JPEG 等图像及其元数据。",
        "标注工作台：支持点、矩形、多边形标注，提供撤销、重做、显隐、删除和对象属性编辑。",
        "AI 辅助：加载带置信度和模型版本的预标注，人工确认后保留修订来源。",
        "任务看板：创建、编辑、删除、筛选并推进标注任务状态。",
        "模型作业：配置模型评估、小样本训练和受控增量训练，并保存运行参数与指标。",
        "数据分析：查看标注产能、标签分布、模型趋势和成员质量。",
    ):
        add_bullet(doc, item)
    add_callout(doc, "使用边界", "本软件首版定位为科研与辅助标注工具，不代替病理医生，不允许未经人工复核的模型结果直接成为金标准。")
    page_break(doc)

    # Page 5
    add_page_title(doc, "第2章", "运行环境与安装", "推荐在 Windows 10/11 工作站或机构内网服务器上运行。")
    doc.add_heading("2.1 建议环境", level=2)
    add_metadata_table(
        doc,
        [
            ("操作系统", "Windows 10/11 64 位；后端也可部署于 Linux"),
            ("浏览器", "Chrome、Edge 等现代浏览器"),
            ("Node.js", "建议 20 LTS 或兼容版本"),
            ("Python", "建议 3.10 及以上"),
            ("内存", "开发演示不少于 8 GB；处理大型 WSI 建议 16 GB 以上"),
            ("磁盘", "按切片规模预留空间；原始 WSI 默认只读"),
        ],
    )
    doc.add_heading("2.2 安装步骤", level=2)
    add_numbered(doc, "打开 PowerShell，进入 PathFISA 项目根目录。")
    add_numbered(doc, "执行 .\\scripts\\setup.ps1，创建 Python 虚拟环境并安装前后端依赖。")
    add_numbered(doc, "安装完成后执行 npm run dev 启动 Web 与 API 服务。")
    add_callout(doc, "依赖说明", "后端使用 FastAPI、Pillow、OpenSlide；前端使用 React、TypeScript、Vite 和 OpenSeadragon。")
    page_break(doc)

    # Page 6
    add_page_title(doc, "第3章", "启动与访问")
    doc.add_heading("3.1 启动系统", level=2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(".\\scripts\\setup.ps1\nnpm run dev")
    set_run_font(r, latin="Consolas", east_asia="宋体", size=10, color=WHITE)
    set_cell_like = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), NAVY)
    set_cell_like.append(shd)
    doc.add_heading("3.2 访问地址", level=2)
    add_metadata_table(
        doc,
        [
            ("Web 页面", "http://127.0.0.1:5173"),
            ("API 文档", "http://127.0.0.1:8000/docs"),
            ("健康检查", "http://127.0.0.1:8000/api/v1/health"),
        ],
    )
    doc.add_heading("3.3 正常启动标志", level=2)
    add_bullet(doc, "PowerShell 中同时显示 web 与 api 两个服务的启动日志。")
    add_bullet(doc, "浏览器可打开总览页面，左侧菜单可进入切片、任务、模型、分析和设置。")
    add_bullet(doc, "API 健康检查返回 status=ok。")
    add_callout(doc, "停止服务", "在启动服务的 PowerShell 窗口中按 Ctrl+C。")
    page_break(doc)

    # Page 7
    add_page_title(doc, "第4章", "主界面与导航", "总览页集中展示项目进度、待处理任务、AI 建议和最近任务。")
    add_picture(doc, "dashboard.png", width_cm=16.6, caption="图4-1 总览页面")
    add_bullet(doc, "左侧导航：总览、病例与切片、标注任务、模型中心、数据分析、项目设置。")
    add_bullet(doc, "顶部搜索：用于查找病例、切片或任务；右上角可打开通知中心和智能助手。")
    add_bullet(doc, "总览卡片：显示切片总数、完成率、待复核任务和 AI 标注接受率。")
    page_break(doc)

    # Page 8
    add_page_title(doc, "第5章", "通知中心与智能助手")
    add_two_pictures(
        doc,
        ("notification-center.png", "图5-1 通知中心"),
        ("intelligent-assistant.png", "图5-2 智能助手"),
    )
    doc.add_heading("5.1 通知中心", level=2)
    add_bullet(doc, "点击顶部铃铛查看任务、复核和模型作业相关通知。")
    add_bullet(doc, "支持标记全部已读，并可从通知跳转到关联页面。")
    doc.add_heading("5.2 智能助手", level=2)
    add_bullet(doc, "可询问标注操作、真实 SVS、任务状态和模型流程。")
    add_bullet(doc, "当前版本使用模拟问答接口，回答用于操作提示，不构成诊断意见。")
    page_break(doc)

    # Page 9
    add_page_title(doc, "第6章", "病例与切片管理")
    add_picture(doc, "real-slides.png", width_cm=16.6, caption="图6-1 病例与切片库")
    add_bullet(doc, "切片卡片展示病例编号、文件名、染色、组织类型、尺寸、倍率和处理状态。")
    add_bullet(doc, "可使用搜索框、染色类型和更多筛选条件定位切片。")
    add_bullet(doc, "点击切片卡片进入标注工作台；原始 SVS 文件按只读方式加载。")
    page_break(doc)

    # Page 10
    add_page_title(doc, "第7章", "切片导入", "系统支持常见图像及病理全切片格式；大文件上传时间取决于磁盘和网络。")
    add_picture(doc, "slides.png", width_cm=16.3, caption="图7-1 切片库与添加入口")
    add_numbered(doc, "在“病例与切片”页面点击“添加切片”或“批量导入”。")
    add_numbered(doc, "选择本地文件，并填写病例编号、染色类型、组织类型、MPP 和物镜倍率。")
    add_numbered(doc, "提交后等待后端校验图像格式和尺寸；成功后切片出现在列表中。")
    add_callout(doc, "支持格式", "JPEG、PNG、TIFF、SVS、NDPI、MRXS、SCN、VMS、VMU、BIF 等；实际可读性受本机 OpenSlide 支持范围影响。")
    page_break(doc)

    # Page 11
    add_page_title(doc, "第8章", "标注工作台", "工作台由顶部信息栏、左侧工具栏、中央 WSI 视图、右侧对象/属性面板和导航缩略图组成。")
    add_picture(doc, "annotation.png", width_cm=16.6, caption="图8-1 标注工作台")
    add_bullet(doc, "顶部显示切片名称、倍率、MPP、云端草稿状态及提交复核按钮。")
    add_bullet(doc, "左侧工具栏用于选择、平移、绘制、测量、撤销、重做和删除。")
    add_bullet(doc, "右侧面板用于图层显隐、对象选择、标签切换、属性编辑和 AI 建议确认。")
    page_break(doc)

    # Page 12
    add_page_title(doc, "第9章", "浏览与显示参数")
    add_picture(doc, "annotation-tools.png", width_cm=16.3, caption="图9-1 显示参数与测量工具")
    add_bullet(doc, "平移：选择手形工具后拖动切片，或使用鼠标滚轮/触控板缩放。")
    add_bullet(doc, "缩放：使用左下角缩放控件，导航缩略图中的视框用于快速定位。")
    add_bullet(doc, "显示参数：调整亮度、对比度和饱和度，仅影响浏览显示，不修改原始图像。")
    add_bullet(doc, "距离测量：选择测量工具后依次点击起点和终点，结果依据 MPP 换算。")
    page_break(doc)

    # Page 13
    add_page_title(doc, "第10章", "创建标注与快捷键")
    doc.add_heading("10.1 绘制标注", level=2)
    add_numbered(doc, "在左上方选择当前标签，确认颜色和类别。")
    add_numbered(doc, "选择点、矩形或多边形工具，在切片视图中绘制。")
    add_numbered(doc, "多边形可通过双击或执行完成动作闭合；新对象自动进入人工标注图层。")
    add_numbered(doc, "选中对象后可移动、查看属性或删除；使用撤销/重做恢复操作。")
    doc.add_heading("10.2 常用快捷键", level=2)
    table = doc.add_table(rows=1, cols=4)
    set_table_widths(table, [1200, 3480, 1200, 3480], indent_dxa=120)
    set_table_borders(table)
    for index, text in enumerate(("按键", "功能", "按键", "功能")):
        set_cell_shading(table.rows[0].cells[index], MINT)
        r = table.rows[0].cells[index].paragraphs[0].add_run(text)
        set_run_font(r, size=9, color=NAVY, bold=True)
    shortcuts = [
        ("V", "选择工具", "H", "平移切片"),
        ("P", "多边形", "R", "矩形"),
        ("D", "点标注", "M", "距离测量"),
        ("Ctrl+S", "保存标注", "Ctrl+Z", "撤销"),
        ("Delete", "删除选中", "Esc", "取消绘制"),
    ]
    for values in shortcuts:
        cells = table.add_row().cells
        for cell, text in zip(cells, values):
            r = cell.paragraphs[0].add_run(text)
            set_run_font(r, size=9, color=INK, bold=text in {"V", "H", "P", "R", "D", "M", "Ctrl+S", "Ctrl+Z", "Delete", "Esc"})
    add_callout(doc, "提示", "快捷键在标注工作台内生效；输入评论或编辑文本时应先确认焦点位置。")
    page_break(doc)

    # Page 14
    add_page_title(doc, "第11章", "标注属性与 AI 建议")
    add_picture(doc, "annotation-ai-properties.png", width_cm=16.2, caption="图11-1 标注对象属性与 AI 诊断提示")
    add_bullet(doc, "对象属性可修改标签类别、标注类型、线条颜色，并填写形态和复核描述。")
    add_bullet(doc, "AI 原始建议采用虚线显示，面板展示置信度、模型版本和提示信息。")
    add_bullet(doc, "点击“人工确认/接收”后，对象来源标记为 AI 接受或 AI 修订，便于追溯。")
    add_bullet(doc, "AI 建议必须由具备权限的人员复核，不得直接作为最终诊断。")
    page_break(doc)

    # Page 15
    add_page_title(doc, "第12章", "保存、提交与复核")
    doc.add_heading("12.1 保存草稿", level=2)
    add_bullet(doc, "点击右下角“保存”或使用 Ctrl+S，将当前对象保存为切片标注 JSON。")
    add_bullet(doc, "保存成功后页面显示已保存状态；重新打开切片时自动加载已保存标注。")
    doc.add_heading("12.2 提交复核", level=2)
    add_numbered(doc, "检查标注对象数量、标签、边界和 AI 来源。")
    add_numbered(doc, "点击顶部“提交复核”，任务进入待复核状态。")
    add_numbered(doc, "复核人员通过任务看板打开切片，查看人工与 AI 结果，决定通过或返修。")
    doc.add_heading("12.3 协作与分享", level=2)
    add_bullet(doc, "评论入口用于记录切片级复核意见；分享按钮复制当前页面地址。")
    add_bullet(doc, "可切换全屏或收起右侧面板，扩大 WSI 浏览区域。")
    add_callout(doc, "数据位置", "标注结果默认保存到项目 data/annotations 目录，原始 WSI 文件不会被改写。")
    page_break(doc)

    # Page 16
    add_page_title(doc, "第13章", "标注任务管理")
    add_picture(doc, "tasks-crud.png", width_cm=16.4, caption="图13-1 标注任务看板")
    add_bullet(doc, "看板按待开始、标注中、待复核、已完成分列展示。")
    add_bullet(doc, "点击“创建任务”设置病例、切片、任务类型、负责人、复核人、优先级和截止日期。")
    add_bullet(doc, "任务支持详情查看、编辑、删除以及状态推进；筛选栏可按关键字和状态定位。")
    page_break(doc)

    # Page 17
    add_page_title(doc, "第14章", "模型中心与模型作业")
    doc.add_heading("14.1 作业类型", level=2)
    add_metadata_table(
        doc,
        [
            ("模型评估", "在指定数据集上记录阈值、Patch 尺寸、重叠率及 Dice/IoU 等指标"),
            ("小样本训练", "选择基础模型、数据集、冻结策略、增强方法和训练参数"),
            ("增量训练", "组合新增金标准、历史回放、蒸馏约束和回归门禁"),
        ],
    )
    doc.add_heading("14.2 基本操作", level=2)
    add_numbered(doc, "进入“模型中心”，选择新建评估、小样本训练或增量训练。")
    add_numbered(doc, "按向导填写基础配置、算法架构、运行参数和质量门禁。")
    add_numbered(doc, "保存草稿或提交作业；系统记录配置、状态、进度和指标结果。")
    add_callout(doc, "受控发布", "当前版本的模型作业以流程配置和模拟运行接口为主；模型发布必须保留人工审批环节。")
    page_break(doc)

    # Page 18
    add_page_title(doc, "第15章", "增量训练配置")
    add_picture(doc, "incremental-workflow.png", width_cm=16.4, caption="图15-1 增量训练配置向导")
    add_bullet(doc, "基础配置：填写作业名称，选择数据集版本、基础模型和训练策略。")
    add_bullet(doc, "算法架构：选择网络结构、回放方案及防遗忘约束。")
    add_bullet(doc, "运行参数：设置训练轮次、学习率、批大小和资源预算。")
    add_bullet(doc, "质量门禁：配置新任务提升、旧任务最大退化和稳定性阈值。")
    add_bullet(doc, "回归门禁：候选模型与当前模型在固定测试集上比较，未通过不得发布。")
    page_break(doc)

    # Page 19
    add_page_title(doc, "第16章", "数据分析")
    add_picture(doc, "analytics.png", width_cm=16.4, caption="图16-1 数据分析看板")
    add_bullet(doc, "顶部指标显示切片数量、标注对象、完成任务、AI 接受率和中位标注耗时。")
    add_bullet(doc, "近七日标注产能用于观察工作量变化；标签分布用于发现类别不平衡。")
    add_bullet(doc, "模型版本趋势展示评估指标变化；成员效率与质量用于项目管理和复核抽检。")
    add_bullet(doc, "点击导出分析报告可形成后续归档或汇报材料。")
    page_break(doc)

    # Page 20
    add_page_title(doc, "第17章", "项目设置与数据管理")
    doc.add_heading("17.1 项目设置", level=2)
    add_bullet(doc, "基础信息：项目名称、项目编码、说明、机构、负责人。")
    add_bullet(doc, "标签体系：标签名称、颜色及业务分类。")
    add_bullet(doc, "标注流程：默认标注人、复核人、复核模式和自动保存间隔。")
    add_bullet(doc, "AI 配置：默认模型、自动加载开关和置信度阈值。")
    add_bullet(doc, "安全与通知：导出权限、导出审批、审计保留周期等。")
    doc.add_heading("17.2 本地数据目录", level=2)
    add_metadata_table(
        doc,
        [
            ("data/annotations", "按切片保存标注 JSON"),
            ("data/slides", "保存通过上传接口导入的切片文件"),
            ("data/slides.json", "导入切片登记信息"),
            ("data/tasks.json", "标注任务记录"),
            ("data/model-runs.json", "模型作业与指标记录"),
            ("data/project-settings.json", "项目设置"),
        ],
    )
    page_break(doc)

    # Page 21
    add_page_title(doc, "第18章", "常见问题")
    faq = [
        ("页面无法打开", "确认 npm run dev 正在运行；检查 5173 和 8000 端口是否被占用。"),
        ("切片显示空白", "检查原始文件路径是否存在、OpenSlide 是否支持该格式，并查看 API 日志。"),
        ("导入失败", "确认扩展名在支持列表内，文件未损坏，且 data 目录具有写权限。"),
        ("标注未保存", "点击保存并等待成功提示；检查 API 是否可访问及 data/annotations 权限。"),
        ("测量结果不准确", "确认切片 MPP 参数正确；MPP 缺失时只能进行像素级参考测量。"),
        ("AI 建议未出现", "确认 AI 辅助开关已开启；当前版本为演示/模拟建议接口。"),
        ("模型作业不运行", "当前首版使用模拟运行流程；请检查作业状态和配置记录。"),
        ("中文或图标异常", "使用现代浏览器，确认前端依赖安装完整并重新构建。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    set_table_widths(table, [2500, 6860], indent_dxa=120)
    set_table_borders(table)
    for idx, text in enumerate(("问题", "处理方法")):
        set_cell_shading(table.rows[0].cells[idx], MINT)
        r = table.rows[0].cells[idx].paragraphs[0].add_run(text)
        set_run_font(r, size=9.5, color=NAVY, bold=True)
    for question, answer in faq:
        cells = table.add_row().cells
        r = cells[0].paragraphs[0].add_run(question)
        set_run_font(r, size=9, color=NAVY, bold=True)
        r = cells[1].paragraphs[0].add_run(answer)
        set_run_font(r, size=9, color=INK)
        for cell in cells:
            set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
    page_break(doc)

    # Page 22
    add_page_title(doc, "第19章", "安全、合规与使用边界")
    doc.add_heading("19.1 数据安全", level=2)
    add_bullet(doc, "病理切片和标注数据应部署在机构授权的工作站、内网或私有环境。")
    add_bullet(doc, "病例信息应使用去标识化编号；导出前确认脱敏范围和审批要求。")
    add_bullet(doc, "定期备份 data 目录及模型产物；原始 WSI 建议保留只读副本和校验值。")
    doc.add_heading("19.2 使用边界", level=2)
    add_bullet(doc, "系统不自动生成最终临床诊断或病理报告。")
    add_bullet(doc, "AI 预标注、智能助手回答和模拟模型指标仅供科研及操作辅助。")
    add_bullet(doc, "未经人工复核的数据不得作为金标准；未经门禁和审批的模型不得发布。")
    add_bullet(doc, "正式进入临床环境前，应完成数据安全、网络安全、个人信息保护、伦理和医疗器械适用性评审。")
    add_callout(doc, "责任提示", "软件使用单位应根据适用法律法规和内部制度配置权限、审计、备份和审批流程。")
    page_break(doc)

    # Page 23
    add_page_title(doc, "附录", "接口、目录与验收清单")
    doc.add_heading("A.1 主要 API", level=2)
    doc.add_paragraph(
        "/api/v1/slides、/api/v1/tasks、/api/v1/model-runs、/api/v1/project-settings、"
        "/api/v1/analytics、/api/v1/notifications、/api/v1/assistant/chat。"
    )
    doc.add_heading("A.2 基本验收清单", level=2)
    for item in (
        "能够启动 Web 与 API，并打开总览页面。",
        "能够查看切片库并打开真实或导入的 WSI。",
        "能够创建点、矩形、多边形标注，并执行保存、撤销、重做和删除。",
        "能够加载 AI 建议、查看置信度并进行人工确认。",
        "能够创建、编辑、删除和推进标注任务。",
        "能够创建模型作业并查看状态、配置和指标。",
        "能够查看数据分析和修改项目设置。",
        "能够在重新打开切片后加载已保存标注。",
    ):
        add_bullet(doc, item)
    doc.add_heading("A.3 材料提交前检查", level=2)
    add_bullet(doc, "将封面中的著作权人和编制单位替换为实际申请主体。")
    add_bullet(doc, "确认软件全称、简称和版本号与申请表完全一致。")
    add_bullet(doc, "确认截图不含真实患者身份信息或其他不宜公开内容。")
    add_bullet(doc, "结合登记机构要求打印或转换为 PDF，并保留 Word 原稿。")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(MANUAL_PATH)


@dataclass
class SourceLine:
    display: str
    is_code: bool


FILE_DESCRIPTIONS = {
    "apps/api/app/main.py": "FastAPI 后端主程序，提供切片、标注、任务、模型、设置与分析接口",
    "apps/api/app/imaging.py": "病理图像读取、缩略图和 Deep Zoom 瓦片处理",
    "apps/api/app/models.py": "Pydantic 数据模型与接口参数定义",
    "apps/api/app/storage.py": "本地 JSON 数据及标注持久化",
    "apps/web/src/api.ts": "前端 API 请求封装",
    "apps/web/src/components/AppShell.tsx": "应用主框架、导航、通知中心和智能助手",
    "apps/web/src/components/Modal.tsx": "通用模态对话框组件",
    "apps/web/src/components/PageHeader.tsx": "页面标题与操作区域组件",
    "apps/web/src/components/SlideViewer.tsx": "WSI 浏览、绘制、属性编辑与 AI 标注核心组件",
    "apps/web/src/pages/AnalyticsPage.tsx": "数据分析页面",
    "apps/web/src/pages/AnnotationPage.tsx": "标注工作台页面",
    "apps/web/src/pages/DashboardPage.tsx": "项目总览页面",
    "apps/web/src/pages/ModelsPage.tsx": "模型中心页面",
    "apps/web/src/pages/ModelWorkflowPage.tsx": "评估、小样本训练与增量训练配置向导",
    "apps/web/src/pages/SettingsPage.tsx": "项目设置页面",
    "apps/web/src/pages/SlidesPage.tsx": "病例与切片管理页面",
    "apps/web/src/pages/TasksPage.tsx": "标注任务增删改查与状态看板",
    "apps/web/src/types.ts": "前端领域类型定义",
    "apps/web/vite.config.ts": "Vite 开发服务器与代理配置",
    "apps/web/src/App.tsx": "前端路由入口",
    "apps/web/src/main.tsx": "React 应用挂载入口",
    "apps/web/src/components/StatusPill.tsx": "状态标签组件",
    "apps/web/src/vite-env.d.ts": "Vite 类型声明",
    "apps/web/vite.config.d.ts": "Vite 配置类型声明",
}


SOURCE_ORDER = [
    "apps/api/app/main.py",
    "apps/api/app/imaging.py",
    "apps/api/app/models.py",
    "apps/api/app/storage.py",
    "apps/web/src/api.ts",
    "apps/web/src/components/AppShell.tsx",
    "apps/web/src/components/Modal.tsx",
    "apps/web/src/components/PageHeader.tsx",
    "apps/web/src/components/SlideViewer.tsx",
    "apps/web/src/pages/AnalyticsPage.tsx",
    "apps/web/src/pages/AnnotationPage.tsx",
    "apps/web/src/pages/DashboardPage.tsx",
    "apps/web/src/pages/ModelsPage.tsx",
    "apps/web/src/pages/ModelWorkflowPage.tsx",
    "apps/web/src/pages/SettingsPage.tsx",
    "apps/web/src/pages/SlidesPage.tsx",
    "apps/web/src/pages/TasksPage.tsx",
    "apps/web/src/types.ts",
    "apps/web/vite.config.ts",
]

RESERVED_LAST_PAGE = [
    "apps/web/src/App.tsx",
    "apps/web/src/main.tsx",
    "apps/web/src/components/StatusPill.tsx",
    "apps/web/src/vite-env.d.ts",
    "apps/web/vite.config.d.ts",
]


def source_file_header(path: str) -> list[SourceLine]:
    language = "Python" if path.endswith(".py") else "TypeScript/TSX"
    description = FILE_DESCRIPTIONS[path]
    return [
        SourceLine("/*******************************************************************************", False),
        SourceLine(f" * 软件名称：{SOFTWARE_NAME} {VERSION}", False),
        SourceLine(f" * 源文件：{path}", False),
        SourceLine(f" * 编程语言：{language}    著作权人：{COPYRIGHT_OWNER}", False),
        SourceLine(f" * 功能说明：{description}", False),
        SourceLine(" *******************************************************************************/", False),
    ]


def load_source_file(path: str) -> list[SourceLine]:
    text = (ROOT / path).read_text(encoding="utf-8")
    result = source_file_header(path)
    for original_number, raw_line in enumerate(text.splitlines(), start=1):
        if not raw_line.strip():
            continue
        normalized = raw_line.replace("\t", "    ").rstrip()
        result.append(SourceLine(f"{original_number:04d} │ {normalized}", True))
    return result


def partition_lines(lines: list[SourceLine], pages: int) -> list[list[SourceLine]]:
    total = len(lines)
    boundaries = [0]
    header_ranges: list[tuple[int, int]] = []
    index = 0
    while index < total:
        if lines[index].display.startswith("/*****"):
            header_ranges.append((index, min(index + 6, total)))
            index += 6
        else:
            index += 1

    for page_no in range(1, pages):
        target = round(page_no * total / pages)
        for start, end in header_ranges:
            if start < target < end:
                target = start
                break
        min_boundary = boundaries[-1] + 58
        max_boundary = total - (pages - page_no) * 58
        target = max(min_boundary, min(target, max_boundary))
        boundaries.append(target)
    boundaries.append(total)
    return [lines[boundaries[i] : boundaries[i + 1]] for i in range(pages)]


def set_source_geometry(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.05)
    section.bottom_margin = Cm(1.05)
    section.left_margin = Cm(1.05)
    section.right_margin = Cm(0.9)
    section.header_distance = Cm(0.35)
    section.footer_distance = Cm(0.35)


def add_source_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{SOFTWARE_NAME} {VERSION} 源程序")
    set_run_font(r, latin="Arial", east_asia="宋体", size=7.2, color=INK, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("第 ")
    set_run_font(r, size=7.2, color=GRAY)
    add_field(p, "PAGE")
    r = p.add_run(" 页 / 共 ")
    set_run_font(r, size=7.2, color=GRAY)
    add_field(p, "NUMPAGES")
    r = p.add_run(" 页")
    set_run_font(r, size=7.2, color=GRAY)


def add_source_page(doc: Document, lines: list[SourceLine], page_number: int) -> None:
    for item in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(7.05)
        p.paragraph_format.keep_together = False
        p.paragraph_format.keep_with_next = False
        p.paragraph_format.widow_control = False
        r = p.add_run(item.display)
        if item.is_code:
            set_run_font(r, latin="Consolas", east_asia="宋体", size=5.8, color="111111")
        else:
            set_run_font(r, latin="Consolas", east_asia="宋体", size=5.8, color=NAVY, bold=True)
    if page_number < 60:
        doc.add_page_break()


def build_source_document() -> None:
    leading: list[SourceLine] = []
    for path in SOURCE_ORDER:
        leading.extend(load_source_file(path))
    last_page: list[SourceLine] = []
    for path in RESERVED_LAST_PAGE:
        last_page.extend(load_source_file(path))

    pages = partition_lines(leading, 59)
    pages.append(last_page)

    for page_index, page in enumerate(pages, start=1):
        code_count = sum(1 for item in page if item.is_code)
        if code_count < 50:
            raise RuntimeError(f"Source page {page_index} has only {code_count} code lines")
    if len(pages) != 60:
        raise RuntimeError(f"Expected 60 pages, got {len(pages)}")
    if not pages[0][0].display.startswith("/*****"):
        raise RuntimeError("First page does not start with a complete file header")
    if not pages[-1][0].display.startswith("/*****"):
        raise RuntimeError("Last page does not start with a complete file header")

    doc = Document()
    set_source_geometry(doc)
    add_source_header_footer(doc)
    normal = doc.styles["Normal"]
    normal.font.name = "Consolas"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(5.8)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(7.05)

    doc.core_properties.title = f"{SOFTWARE_NAME}{VERSION}源码文档"
    doc.core_properties.subject = "计算机软件著作权登记源程序鉴别材料"
    doc.core_properties.author = "PathFISA"
    doc.core_properties.keywords = "PathFISA, 源程序, 软件著作权, Python, TypeScript"

    for page_number, page in enumerate(pages, start=1):
        add_source_page(doc, page, page_number)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(SOURCE_PATH)

    audit_path = OUTPUT_DIR / "source_page_audit.txt"
    audit_lines = [
        f"page={index:02d}, code_lines={sum(1 for item in page if item.is_code)}, display_lines={len(page)}"
        for index, page in enumerate(pages, start=1)
    ]
    audit_path.write_text("\n".join(audit_lines), encoding="utf-8")


def main() -> None:
    build_manual()
    build_source_document()
    print(MANUAL_PATH)
    print(SOURCE_PATH)


if __name__ == "__main__":
    main()
